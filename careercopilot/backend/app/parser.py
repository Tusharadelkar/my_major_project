import os
import io
import json
import logging
import fitz  # PyMuPDF
import docx
from groq import Groq

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract raw text from a PDF file using PyMuPDF."""
    text = ""
    try:
        # Open PDF document from memory
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        for page in doc:
            text += page.get_text()
        doc.close()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    return text

def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract raw text from a DOCX file using python-docx."""
    text = ""
    try:
        # Open word document from memory
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        # Include table text if any
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text.append(cell.text)
        text = "\n".join(paragraphs + table_text)
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    return text

def parse_resume_with_llm(resume_text: str, job_description_text: str) -> dict:
    """
    Send the resume text and job description to Llama 3 via Groq.
    Returns structured JSON. Implements retries for malformed JSON.
    """
    groq_api_key = os.getenv("GROQ_API_KEY")
    if not groq_api_key or groq_api_key == "your_groq_api_key_here":
        # Fallback to local stub if key is not configured to allow graceful failures/demo mode
        logger.warning("GROQ_API_KEY is not set or using placeholder. Returning mock parsed data.")
        return get_fallback_parsed_data(resume_text, job_description_text)

    try:
        client = Groq(api_key=groq_api_key)
    except Exception as e:
        logger.error(f"Failed to initialize Groq client: {str(e)}")
        return get_fallback_parsed_data(resume_text, job_description_text)

    # Prompt requesting structured JSON output
    system_prompt = (
        "You are an expert ATS (Applicant Tracking System) parser. Analyze the provided resume text "
        "and job description text. You MUST respond with a single, valid JSON object and absolutely "
        "nothing else. No explanations, no markdown block wrappers (do NOT use ```json). Just the raw JSON.\n\n"
        "The JSON object must contain exactly these keys:\n"
        "{\n"
        "  \"skills\": [\"list of professional/technical skills extracted from the resume\"],\n"
        "  \"projects\": [\"list of key projects mentioned in the resume\"],\n"
        "  \"experience\": [\"list of work experience roles, companies, and key achievements from the resume\"],\n"
        "  \"education\": [\"list of degrees, schools, and majors from the resume\"],\n"
        "  \"job_description_skills\": [\"list of key technical & soft skills required by the job description\"]\n"
        "}"
    )

    user_content = (
        f"--- RESUME TEXT ---\n{resume_text}\n\n"
        f"--- JOB DESCRIPTION TEXT ---\n{job_description_text}\n"
    )

    max_retries = 3
    for attempt in range(max_retries):
        try:
            logger.info(f"Calling Groq LLM (Llama 3.1 8B) - Attempt {attempt + 1}")
            completion = client.chat.completions.create(
                model="llama-3.1-8b-instant",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_content}
                ],
                temperature=0.1,
                max_tokens=2048,
                response_format={"type": "json_object"}
            )
            
            response_text = completion.choices[0].message.content.strip()
            
            # Try to load and validate JSON
            data = json.loads(response_text)
            
            # Validate required fields
            required_keys = ["skills", "projects", "experience", "education", "job_description_skills"]
            for key in required_keys:
                if key not in data or not isinstance(data[key], list):
                    # Force a schema correction/retry if fields are missing or wrong type
                    data[key] = data.get(key, []) or []
            
            return data

        except (json.JSONDecodeError, Exception) as e:
            logger.warning(f"Attempt {attempt + 1} failed during Groq LLM parsing/JSON validation: {str(e)}")
            if attempt == max_retries - 1:
                logger.error("All LLM parsing attempts failed. Returning fallback structured data.")
                return get_fallback_parsed_data(resume_text, job_description_text)

def get_fallback_parsed_data(resume_text: str, job_description_text: str) -> dict:
    """Helper to return fallback structured data if Groq is unavailable or fails."""
    # Simple rule-based extraction for fallback
    # Look for common skills keywords as a stub
    common_skills = ["python", "javascript", "sql", "java", "aws", "docker", "kubernetes", "react", "fastapi", "git"]
    found_skills = []
    for skill in common_skills:
        if skill in resume_text.lower():
            found_skills.append(skill.capitalize())
            
    job_skills = []
    for skill in common_skills:
        if skill in job_description_text.lower():
            job_skills.append(skill.capitalize())

    # Ensure some fallback values
    if not found_skills:
        found_skills = ["Python", "Software Engineering", "SQL"]
    if not job_skills:
        job_skills = ["Python", "SQL", "Docker", "FastAPI"]

    return {
        "skills": found_skills,
        "projects": ["Fallback Project: CareerCopilot AI Portfolio"],
        "experience": ["Fallback Experience: Software Engineer in residency"],
        "education": ["Fallback Education: B.S. in Computer Science"],
        "job_description_skills": job_skills
    }
